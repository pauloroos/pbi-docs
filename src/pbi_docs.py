import requests
import os
import json
import subprocess
import time
import psutil
import pandas as pd
import clr
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from pathlib import Path
import re

# Crendentials
tenant_id = os.getenv("FABRIC_TENANT_ID")
client_id = os.getenv("FABRIC_CLIENT_ID")
client_secret = os.getenv("FABRIC_CLIENT_SECRET")

# Path DAX Studio CLI
cmd = r"C:\Program Files\DAX Studio\dscmd.exe"

# Path Analysis Services
ssas_dll = r"C:\Program Files\DAX Studio\bin\Microsoft.AnalysisServices.dll"

# Path Power BI Desktop
pbi_desktop = r"C:\Program Files\WindowsApps\Microsoft.MicrosoftPowerBIDesktop_2.140.1205.0_x64__8wekyb3d8bbwe\bin\PBIDesktop.exe"

# DAX Query 
dax_query = """
DEFINE

// Source Tables
VAR __tables = INFO.TABLES()
VAR __columns = INFO.COLUMNS()
VAR __measures = INFO.MEASURES()
VAR __formats = INFO.FORMATSTRINGDEFINITIONS()
VAR __relationships = INFO.RELATIONSHIPS()
VAR __calculationGroups = INFO.CALCULATIONGROUPS()
VAR __calculationItems = INFO.CALCULATIONITEMS()
VAR __dependencies = INFO.CALCDEPENDENCY()
VAR __partitions = INFO.PARTITIONS()


// Measures
VAR __measuresPre = 
	ADDCOLUMNS(
    		__measures,
    		"Table",
		VAR TableID = [TableID] RETURN
	    	MAXX(FILTER(__tables, [ID] = TableID), [Name]),
    		"Format",
		VAR MeasureID = [ID] RETURN
	    	MAXX(FILTER(__formats, [ObjectID] = MeasureID), [Expression])
	)

VAR __measuresResult =
SELECTCOLUMNS(
	__measuresPre,
	"tableName", [Table],
	"name", [Name],
	"expression", [Expression],
	"format", [FormatString],
	"isHidden", [IsHidden],
	"description", [Description],
	"type",
	SWITCH(
		[DataType],
		2, "String",
		6, "Whole Number",
		8, "Double",
		9, "Datetime",
		10, "Currency",
		11, "Boolean"
	),
	"displayFolder", [DisplayFolder]
) 

// Columns
VAR __columnsPre =
ADDCOLUMNS(
	FILTER(
		__columns,
		NOT CONTAINSSTRING([ExplicitName], "RowNumber")
	),
	"TableName",
	VAR TableID = [TableID] RETURN
	MAXX(FILTER(__tables, [ID] = TableID), [Name])
)

VAR __columnsResult =
SELECTCOLUMNS(
	__columnsPre,
    "tableName", [TableName],
    "name", COALESCE([ExplicitName], [InferredName]),
	"column", 
		"'" & [TableName] & "'" & 
		"[" & COALESCE([ExplicitName], [InferredName]) & "]",
    "sortedBy",
		VAR OrderID = [SortByColumnID] RETURN
	    MAXX(
			FILTER(__columnsPre, [ID] = OrderID), 
			"'" & [TableName] & "'" & 
		"[" & COALESCE([ExplicitName], [InferredName]) & "]"),
    "format", COALESCE([FormatString], "String"),
    "displayFolder", [DisplayFolder],
    "isHidden", [IsHidden],
    "expression", [Expression]
)

// Relationships
VAR relationshipsPre =
ADDCOLUMNS(
    __relationships,
    "FromTable",
		VAR FromTable = [FromTableID] RETURN
		MAXX(FILTER(__columnsPre, [TableID] = FromTable), [TableName]),
    "ToTable",
		VAR ToTable = [ToTableID] RETURN
	    MAXX(FILTER(__columnsPre, [TableID] = ToTable), [TableName]),
    "FromColumn",
		VAR FromColumn = [FromColumnID] RETURN
	    MAXX(
			FILTER( __columnsPre, [ID] = FromColumn ),
			COALESCE( [ExplicitName], [InferredName] )),
    "ToColumn",
	VAR ToColumn = [ToColumnID]
	RETURN
	    MAXX(
		FILTER( __columnsPre, [ID] = ToColumn ),
		COALESCE( [ExplicitName], [InferredName] )
	    )
)

VAR __relationshipsResult =
SELECTCOLUMNS(
    relationshipsPre,
    "from", "'" & [FromTable] & "'" & "[" & [FromColumn] & "]",
	"fromCardinality", IF([FromCardinality] = 2, "*", [FromCardinality]),
	"to",	"'" & [ToTable] & "'" & "[" & [ToColumn] & "]",
    "toCardinality", IF( [ToCardinality] = 2, "*", [ToCardinality]),
    "isActive", [IsActive],
    "isBidirectional", IF([CrossFilteringBehavior] = 2, "True", "False"),
    "relationship",
		VAR fromConcat = "'" & [FromTable] & "'" & "[" & [FromColumn] & "]"
		VAR toConcat = "'" & [ToTable] & "'" & "[" & [ToColumn] & "]"
		VAR fromCardinality = IF([FromCardinality] = 2, "*", [FromCardinality])
		VAR toCardinality = IF([ToCardinality] = 2, "*", [ToCardinality])
		VAR arrow = IF([CrossFilteringBehavior] = 2, "<->", "<-")
		RETURN
	    	fromConcat & " " & fromCardinality & " " & arrow & " " & toCardinality & " " & toConcat
)

// Calculation Groups
VAR __calculationItemsPre =
ADDCOLUMNS(
    ADDCOLUMNS(
	__calculationItems,
	"TableID",
	    VAR CalculationGroup_ID = [CalculationGroupID] RETURN
		MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [TableID]),
	"Precedence",
	    VAR CalculationGroup_ID = [CalculationGroupID] RETURN
		MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [Precedence])
    ),
    "CalculationGroup",
		VAR Table_ID = [TableID] RETURN
	    MAXX(FILTER(__tables, [ID] = Table_ID), [Name]),
    "Format",
		VAR Format_ID = [FormatStringDefinitionID] RETURN
	    MAXX(FILTER(__formats, [ID] = Format_ID), [Expression]),
    "CalculationGroupColumn",
		VAR Table_ID = [TableID]
		VAR Coluna_ID = MINX(FILTER(__columns, [TableID] = Table_ID), [ID])
		RETURN
		    MAXX(
				FILTER( __columns, [ID] = Coluna_ID ),
				COALESCE( [ExplicitName], [InferredName] )
		    ),
    "OrdinalColumn",
		VAR Table_ID = [TableID]
		VAR Coluna_ID = MAXX(FILTER(__columns, [TableID] = Table_ID), [ID])
		RETURN
		    MAXX(
				FILTER( __columns, [ID] = Coluna_ID ),
				COALESCE( [ExplicitName], [InferredName] )
		    )
)

VAR __calculationGroupsResult =
SELECTCOLUMNS(
    __calculationItemsPre,
    "calculationGroup", [CalculationGroup],
    "precedence", [Precedence],
    "calculationItem", [Name],
    "expression", [Expression],
    "format", [Format],
    "ordinal", [Ordinal],
    "calculationGroupColumn", [CalculationGroupColumn],
    "ordinalColumn", [OrdinalColumn]
)

VAR __parametersResult =
SELECTCOLUMNS(
	SUMMARIZE(
	    FILTER(
			__dependencies,
			[REFERENCED_OBJECT_TYPE] = "M_EXPRESSION" && 
			CONTAINSSTRING( [REFERENCED_EXPRESSION], "IsParameterQuery" )
		),
		[REFERENCED_OBJECT],
		[REFERENCED_EXPRESSION]
	),
    "name", [REFERENCED_OBJECT],
    "expression", [REFERENCED_EXPRESSION]
)


// Queries M or DAX
VAR __partitionsResult =
SELECTCOLUMNS(
    __partitions,
    "tableName", 
	IF(
        NOT(ISBLANK([Name])) && LEN([Name]) > 37,
        LEFT([Name], LEN([Name]) - 37),
        [Name]
    ),
    "description", [Description],
    "queryDefinition", [QueryDefinition],
    "modifiedTime", [ModifiedTime],
    "refreshedTime", [RefreshedTime],
    "type", SWITCH( [Type], 4, "M", 2, "DAX", 7, "Internal" )
)

// Output
EVALUATE __measuresResult ORDER BY [tableName] ASC, [name] ASC
EVALUATE __columnsResult ORDER BY [tableName] ASC, [name] ASC
EVALUATE __relationshipsResult ORDER BY [relationship] ASC
EVALUATE __calculationGroupsResult ORDER BY [precedence] ASC, [ordinal] ASC
EVALUATE __parametersResult 
EVALUATE __partitionsResult ORDER BY [tableName] ASC 
"""

def get_token(tenant_id, client_id, client_secret):

    # Azure Entra Endpoint
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"

    # Requisition data
    token_data = {
        "grant_type": "client_credentials",
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://analysis.windows.net/powerbi/api/.default"
    }

    # Get token
    token_response = requests.post(token_url, data=token_data)
    token_json = token_response.json()

    if "access_token" not in token_json:
        raise Exception(f"Error getting token: {token_json}")

    access_token = token_json["access_token"]
    
    return (access_token)

def get_tenant_metadata(access_token):
    
    if access_token:
        headers = {"Authorization": f"Bearer {access_token}"}

    # Get Workspaces
    workspaces_url = "https://api.powerbi.com/v1.0/myorg/groups"
    workspaces_response = requests.get(workspaces_url, headers=headers)
    workspaces_data = workspaces_response.json().get("value", [])

    if workspaces_response.status_code != 200:
        print("Error getting workspaces:", workspaces_response.text)
        exit()

    print(f"Found {len(workspaces_data)} workspaces.")

    # Separate JSON for workspaces, datasets, reports, and dataflows
    workspaces_info = []
    datasets_info = []
    reports_info = []
    dataflows_info = []

    # Get Workspaces, Datasets, Reports, and Dataflows
    for ws in workspaces_data:
        ws_id = ws["id"]
        ws_name = ws["name"]

        # Save all Workspace information
        workspace_entry = {"workspaceId": ws_id}
        workspace_entry.update(ws)  # Add all workspace details
        workspaces_info.append(workspace_entry)

        # Get Datasets from each Workspace
        datasets_url = f"https://api.powerbi.com/v1.0/myorg/groups/{ws_id}/datasets"
        datasets_response = requests.get(datasets_url, headers=headers)

        if datasets_response.status_code == 200:
            datasets_data = datasets_response.json().get("value", [])
            
            if not datasets_data:
                print(f"No datasets found in workspace: {ws_name}")
            
            for ds in datasets_data:
                dataset_entry = {"workspace": ws_name, "workspaceId": ws_id}
                dataset_entry.update(ds) 
                datasets_info.append(dataset_entry)
        else:
            print(f"Error getting datasets for workspace {ws_name}: {datasets_response.text}")

        # Get Reports from each Workspace
        reports_url = f"https://api.powerbi.com/v1.0/myorg/groups/{ws_id}/reports"
        reports_response = requests.get(reports_url, headers=headers)

        if reports_response.status_code == 200:
            reports_data = reports_response.json().get("value", [])
            
            if not reports_data:
                print(f"No reports found in workspace: {ws_name}")
            
            for rp in reports_data:
                report_entry = {"workspace": ws_name, "workspaceId": ws_id}
                report_entry.update(rp)  
                reports_info.append(report_entry)
        else:
            print(f"Error getting reports for workspace {ws_name}: {reports_response.text}")

        # Get Dataflows from each Workspace
        dataflows_url = f"https://api.powerbi.com/v1.0/myorg/groups/{ws_id}/dataflows"
        dataflows_response = requests.get(dataflows_url, headers=headers)

        if dataflows_response.status_code == 200:
            dataflows_data = dataflows_response.json().get("value", [])
            
            if not dataflows_data:
                print(f"No dataflows found in workspace: {ws_name}")
            
            for df in dataflows_data:
                dataflow_entry = {"workspace": ws_name, "workspaceId": ws_id}
                dataflow_entry.update(df)  
                dataflows_info.append(dataflow_entry)
        else:
            print(f"Error getting dataflows for workspace {ws_name}: {dataflows_response.text}")

    # Create output folder if it doesn't exist
    dir = "results/tenant_metadata"
    os.makedirs(dir, exist_ok=True)

    # Save JSONs to files
    with open(os.path.join(dir, "workspaces.json"), "w", encoding="utf-8") as f:
        json.dump(workspaces_info, f, indent=4, ensure_ascii=False)

    with open(os.path.join(dir, "datasets.json"), "w", encoding="utf-8") as f:
        json.dump(datasets_info, f, indent=4, ensure_ascii=False)

    with open(os.path.join(dir, "reports.json"), "w", encoding="utf-8") as f:
        json.dump(reports_info, f, indent=4, ensure_ascii=False)

    with open(os.path.join(dir, "dataflows.json"), "w", encoding="utf-8") as f:
        json.dump(dataflows_info, f, indent=4, ensure_ascii=False)

    print(f"JSONs saved in folder '{dir}':")
    print(f"   - {dir}/workspaces.json")
    print(f"   - {dir}/datasets.json")
    print(f"   - {dir}/reports.json")
    print(f"   - {dir}/dataflows.json")

def get_info_datasets(cmd, tenant_id, client_id, client_secret, dax_query):
    # Load JSON files
    with open('results/tenant_metadata/workspaces.json', 'r') as f:
        workspaces = json.load(f)

    with open('results/tenant_metadata/datasets.json', 'r') as f:
        datasets = json.load(f)

    query = dax_query

    # Iterate through each workspace
    for workspace in workspaces:
        workspace_name = workspace['name']
        print(f"Processing workspace: {workspace_name}")
        
        # Filter datasets for current workspace
        workspace_datasets = [ds for ds in datasets if ds['workspace'] == workspace_name]
        
        # Iterate through each dataset in the workspace
        for dataset in workspace_datasets:
            dataset_name = dataset['name']
            print(f"  Processing dataset: {dataset_name}")
            
            # Connection String
            server = f"powerbi://api.powerbi.com/v1.0/myorg/{workspace_name}"
            database = dataset_name

            # Output path
            output_dir = f"results/datasets_info/{workspace_name}/{dataset_name}"
            os.makedirs(output_dir, exist_ok=True)
            output_path = f"{output_dir}/extract.csv"

            # Clean up any existing extract files before execution
            for old_name in ["extract.csv", "extract_2.csv", "extract_3.csv", 
                            "extract_4.csv", "extract_5.csv", "extract_6.csv"]:
                old_path = os.path.join(output_dir, old_name)
                if os.path.exists(old_path):
                    os.remove(old_path)
                    print(f"    Removed existing source file: {old_name}")

            # Executing
            subprocess.run([
                cmd, "csv", output_path,
                "-s", server,
                "-d", database,
                "-u", f"app:{client_id}@{tenant_id}",
                "-p", client_secret,
                "-q", query,
                "-n"
            ])

            # Remaping the extracted files
            rename_map = {
                "extract.csv": "measures.csv",
                "extract_2.csv": "columns.csv",
                "extract_3.csv": "relationships.csv",
                "extract_4.csv": "calculation_groups.csv",
                "extract_5.csv": "parameters.csv",
                "extract_6.csv": "partitions.csv"
            }

            # Rename files
            for old_name, new_name in rename_map.items():
                old_path = os.path.join(output_dir, old_name)
                new_path = os.path.join(output_dir, new_name)

                # Check if old file exists
                if os.path.exists(old_path):
                    # If new file already exists, remove it first
                    if os.path.exists(new_path):
                        os.remove(new_path)
                        print(f"    Removed existing file: {new_name}")
                    
                    os.rename(old_path, new_path)
                    print(f"    Renamed: {old_name} -> {new_name}")
                else:
                    print(f"    File not found: {old_name}")

    print("Processing completed for all workspaces and datasets.")

def get_dataflows(access_token):
    if access_token:
        headers = {"Authorization": f"Bearer {access_token}"}

    # Load daflows.json file
    with open('results/tenant_metadata/dataflows.json', 'r') as f:
        dataflows = json.load(f)

    for dataflow in dataflows:
        workspace_id = dataflow['workspaceId']
        workspace_name =  dataflow['workspace']
        dataflow_id =  dataflow['objectId']
        dataflow_name =  dataflow['name']
        print(f"Processing dataflow [{dataflow_name}] from workspace [{workspace_name}]")

        # Get Dataflows
        dataflow_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/dataflows/{dataflow_id}"

        dataflow_response = requests.get(dataflow_url, headers=headers)
        dataflow_response.raise_for_status()

        # Output path
        output_dir = "results/dataflows_json"
        os.makedirs(output_dir, exist_ok=True)

        dataflow_json = dataflow_response.json()
        with open(f"{output_dir}/{workspace_name}${dataflow_name}.json", "w", encoding="utf-8") as file:
            json.dump(dataflow_json, file, indent=4, ensure_ascii=False)

def get_pro_datasets(access_token):
            
    if access_token:
        headers = {"Authorization": f"Bearer {access_token}"}
    
    # Load workspaces.json file
    with open('results/tenant_metadata/workspaces.json', 'r') as f:
        workspaces = json.load(f)

    # Filtering workspaces IDs where isOnDedicatedCapacity=False
    filtered_workspace_ids = [ws["id"] for ws in workspaces if not ws.get("isOnDedicatedCapacity")]

    # Load reports.json file
    with open('results/tenant_metadata/reports.json', 'r') as f:
        reports = json.load(f)

    # Filtering reports where workspace ID is in filtered_workspaces_ids
    filtered_reports = [report for report in reports if report["workspaceId"] in filtered_workspace_ids]

    # Iterate through each workspace
    for report in filtered_reports:
        # Workspace and Report Info
        workspace_id = report['workspaceId']
        workspace_name =  report['workspace']
        report_id =  report['id']
        report_name =  report['name']
        print(f"Processing report [{report_name}] from workspace [{workspace_name}]")

        # Endpoint
        export_url = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/reports/{report_id}/Export"
        export_response = requests.get(export_url, headers=headers, stream=True)

        # Output path
        output_dir = f"results/exported_pbix/{workspace_name}"
        os.makedirs(output_dir, exist_ok=True)

        # Save the exported PBIX
        if export_response.status_code == 200:
            pbix_path = f"{output_dir}/{report_name}.pbix"
            with open(pbix_path, "wb") as file:
                for chunk in export_response.iter_content(chunk_size=1024):
                    file.write(chunk)
            
            print(f"PBIX saved as: {pbix_path}")
        elif export_response.status_code == 404:
            print("[404]: Workspace or report not found. Check IDs and permissions.")
        elif export_response.status_code == 400:
            print(f"[400]: Invalid request: {export_response.text}")
        else:
            print(f"Inspered error [{export_response.status_code}]: {export_response.text}")

def get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop):
    
    # Load SSAS assembly
    clr.AddReference(ssas_dll)
    from Microsoft.AnalysisServices import Server 

    query = dax_query

    # Wait for Power BI Open
    def wait_for_powerbi_to_open(timeout=60, check_interval=5):
        """
        Waits for Power BI Desktop to open and start the Analysis Services server.

        timeout: Maximum wait time in seconds.
        check_interval: Interval between checks in seconds.
        """
        elapsed_time = 0
        while elapsed_time < timeout:
            server_name = get_powerbi_port()
            if server_name != "No instance of Power BI Desktop found.":
                print(f"Power BI Desktop started! Server found at: {server_name}")
                return server_name
            print(f"Waiting for Power BI Desktop to open... ({elapsed_time}s)")
            time.sleep(check_interval)
            elapsed_time += check_interval

        print("Timeout reached. Power BI Desktop did not start correctly.")
        return None

    def get_powerbi_port():
        # Search the process msmdsrv.exe (Analysis Services Power BI)
        for proc in psutil.process_iter(attrs=['pid', 'name']):
            if proc.info['name'] == "msmdsrv.exe":
                pid = proc.info['pid']
                
                # Run netstat to capture active conections
                result = subprocess.run(["netstat", "-ano"], capture_output=True, text=True)
                connections = result.stdout.splitlines()
                
                # Filter corresponding line to PID of msmdsrv.exe
                for line in connections:
                    if str(pid) in line and "LISTENING" in line:
                        match = re.search(r':(\d+)', line)
                        if match:
                            return f"localhost:{match.group(1)}"
        
        return "No instance of Power BI Desktop found."

    # Function to get the database ID
    def get_ssas_database_id(server_name):
        
        server = Server()
        try:
            server.Connect(f"Data Source={server_name}")
            # Assuming there is only one active database
            if server.Databases.Count > 0:
                db_id = server.Databases[0].ID
                return db_id
            else:
                return None
        finally:
            server.Disconnect()

    def wait_for_powerbi_to_close(timeout=30, check_interval=2):
        """
        Waits for Power BI Desktop to close completely.
        """
        elapsed_time = 0
        while elapsed_time < timeout:
            if "PBIDesktop.exe" not in (p.name() for p in psutil.process_iter()):
                print("Power BI Desktop closed successfully.")
                return True
            print(f"Waiting for Power BI Desktop to close... ({elapsed_time}s)")
            time.sleep(check_interval)
            elapsed_time += check_interval
        print("Timeout waiting for Power BI Desktop to close.")
        return False

    # Load workspaces.json file
    with open('results/tenant_metadata/workspaces.json', 'r') as f:
        workspaces = json.load(f)

    # Filtering workspaces IDs where isOnDedicatedCapacity=False
    filtered_workspace_ids = [ws["id"] for ws in workspaces if not ws.get("isOnDedicatedCapacity")]

    # Load reports.json file
    with open('results/tenant_metadata/reports.json', 'r') as f:
        reports = json.load(f)

    # Filtering reports where workspace ID is in filtered_workspaces_ids
    filtered_reports = [report for report in reports if report["workspaceId"] in filtered_workspace_ids]

    # Iterate through each workspace
    for report in filtered_reports:
        # Workspace and Report Info
        workspace_id = report['workspaceId']
        workspace_name =  report['workspace']
        report_id =  report['id']
        report_name =  report['name']
        print(f"Processing report [{report_name}] from workspace [{workspace_name}]")

        pbix_file = os.path.abspath(f"results/exported_pbix/{workspace_name}/{report_name}.pbix")

        # Execute
        if not os.path.exists(pbix_file):
            print(f"Error: The file {pbix_file} was not found.")
        else:
            try:
                subprocess.Popen([pbi_desktop, pbix_file])
                print(f"Power BI Desktop opened the file {pbix_file}.")
            except Exception as e:
                print(f"Error opening Power BI Desktop: {e}")

        # Get Power BI Desktop Server
        server_name = wait_for_powerbi_to_open()

        if not server_name:
            print("Failed to start Power BI. Retrying...")
            time.sleep(5)
            server_name = wait_for_powerbi_to_open()

        if server_name:
            print(f"Server: {server_name}")
            database_id = get_ssas_database_id(server_name=server_name)
            if database_id:
                print(f"Database: {database_id}")
            else:
                print("No database found. Retrying after waiting...")
                time.sleep(5)
                database_id = get_ssas_database_id(server_name=server_name)
                if not database_id:
                    print("Error: Unable to retrieve database ID.")
                    continue  # Skip this, try next
        else:
            print("No server found. Skipping this report.")
            continue

        # Dax Studio Variables
        server = server_name
        database = database_id

        # Output path
        output_dir = f"results/datasets_info/{workspace_name}/{report_name}"
        os.makedirs(output_dir, exist_ok=True)
        output_path = f"{output_dir}/extract.csv"

        # Clean up any existing extract files before execution
        for old_name in ["extract.csv", "extract_2.csv", "extract_3.csv", 
                        "extract_4.csv", "extract_5.csv", "extract_6.csv"]:
            old_path = os.path.join(output_dir, old_name)
            if os.path.exists(old_path):
                os.remove(old_path)
                print(f"    Removed existing source file: {old_name}")

        max_retries = 3
        retry_delay = 5

        for attempt in range(max_retries):
            try:
                # Executing
                subprocess.run([
                    cmd, "csv", output_path,
                    "-s", server,
                    "-d", database,
                    "-q", query
                ], check=True)
                break
            except subprocess.CalledProcessError:
                print(f"Error executing dscmd. Attempt {attempt + 1} of {max_retries}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    print("Persistent failure executing dscmd. Skipping this report.")


        # Close Power BI Desktop
        subprocess.run(["taskkill", "/F", "/IM", "PBIDesktop.exe"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print("Power BI Desktop was closed.")
        
        
        # Remaping the extracted files
        rename_map = {
            "extract.csv": "measures.csv",
            "extract_2.csv": "columns.csv",
            "extract_3.csv": "relationships.csv",
            "extract_4.csv": "calculation_groups.csv",
            "extract_5.csv": "parameters.csv",
            "extract_6.csv": "partitions.csv"
        }

        # Rename files
        for old_name, new_name in rename_map.items():
            old_path = os.path.join(output_dir, old_name)
            new_path = os.path.join(output_dir, new_name)

            # Check if old file exists
            if os.path.exists(old_path):
                # If new file already exists, remove it first
                if os.path.exists(new_path):
                    os.remove(new_path)
                    print(f"    Removed existing file: {new_name}")
            
                os.rename(old_path, new_path)
                print(f"    Renamed: {old_name} -> {new_name}")
            else:
                print(f"    File not found: {old_name}")

        
        # Wait for the process to actually terminate before opening the next file
        time.sleep(5)
        wait_for_powerbi_to_close()


    print("Processing completed for all workspaces with no dedicated capacity.")

def create_documentation():
    def create_semantic_model_doc(workspace_name, dataset_name):
        
        directory = f"results/datasets_info/{workspace_name}/{dataset_name}"
        doc = Document()
        
        # Ajustar as margens do documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # Custom Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(3)

        # Main Title
        doc.add_heading("Semantic Model Technical Documentation", level=1)
        
        # Overview
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(f"Workspace: {workspace_name}")
        doc.add_paragraph(f"Semantic Model: {dataset_name}")
        doc.add_paragraph("This document describes the extracted semantic model structure, including tables, columns, relationships, partitions, measures, calculation groups, and parameters.")
        
        # Load csv files
        file_paths = {
            "Calculation Groups": os.path.join(directory, "calculation_groups.csv"),
            "Columns": os.path.join(directory, "columns.csv"),
            "Measures": os.path.join(directory, "measures.csv"),
            "Parameters": os.path.join(directory, "parameters.csv"),
            "Partitions": os.path.join(directory, "partitions.csv"),
            "Relationships": os.path.join(directory, "relationships.csv"),
        }
        
        dataframes = {name: pd.read_csv(path, encoding="utf-8", delimiter=";") for name, path in file_paths.items() if os.path.exists(path)}

        # Nan to ""
        for key in dataframes:
            dataframes[key] = dataframes[key].fillna("")

        # Apply style tables
        def apply_table_style(table):
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.paragraphs:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i % 2 == 0:
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F3F3F3"/>'.format(nsdecls('w'))))

        # Add sections
        sections = {
            "Columns": ["Columns", ["Table", "Column", "Sorted By", "Format", "Display Folder", "Hidden", "Expression"], ["tableName","name","sortedBy","format","displayFolder","isHidden","expression"]],
            "Partitions": ["Partitions", ["Table", "Type", "Query"], ["tableName", "type", "queryDefinition"]],
            "Relationships": ["Relationships", ["Relationship", "isActive", "isBidirectional"], ["relationship", "isActive", "isBidirectional"]],
            "Measures": ["Measures", ["Name", "Expression", "Description", "Format"], ["name", "expression", "description", "format"]],
            "Calculation Groups": ["Calculation Groups", ["Group", "Item", "Expression", "Ordinal", "Format"], ["calculationGroup", "calculationItem", "expression", "ordinal", "format"]],
            "Parameters": ["Parameters", ["Name", "Expression"], ["name", "expression"]],
        }
        
        for key, (title, headers, columns) in sections.items():
            if key in dataframes:
                doc.add_heading(title, level=2)
                table = doc.add_table(rows=1, cols=len(headers))

                # Set column widths

                if key == "Partitions":
                    table.columns[0].width = Cm(1)  # Table
                    table.columns[1].width = Cm(1)  # Type
                    table.columns[2].width = Cm(48)  # Query

                elif key == "Relationships":
                    table.columns[0].width = Cm(24)  # Relationship
                    table.columns[1].width = Cm(1)  # isActive
                    table.columns[2].width = Cm(1)  # isBidirectional

                elif key == "Measures":
                    table.columns[0].width = Cm(2)  # Name
                    table.columns[1].width = Cm(36)  # Expression
                    table.columns[2].width = Cm(36)  # Description
                    table.columns[3].width = Cm(2)  # Format

                elif key == "Calculation Groups":
                    table.columns[0].width = Cm(2)  # Group
                    table.columns[1].width = Cm(2)  # Item
                    table.columns[2].width = Cm(36)  # Expression
                    table.columns[3].width = Cm(2)  # Ordinal
                    table.columns[4].width = Cm(8)  # Format

                elif key == "Parameters":
                    table.columns[0].width = Cm(2)  # Name
                    table.columns[1].width = Cm(36)  # Expression

                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for _, row in dataframes[key].iterrows():
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        if i < len(row_cells):
                            row_cells[i].text = str(row[col]) if col in row else ""
                
                apply_table_style(table)
        
        doc.add_paragraph("---")
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Document generated on: {now}")
        
        # Ensure the directory exists
        dir_path = "results/documentation"
        os.makedirs(dir_path, exist_ok=True)
        
        file_path = f"{dir_path}/{workspace_name}${dataset_name}.docx"
        doc.save(file_path)
        print(f"File {file_path} saved.")


    # Load reports.json file
    with open('results/tenant_metadata/reports.json', 'r') as f:
        reports = json.load(f)

    # Iterate through each workspace
    for report in reports:
        # Workspace and Report Info
        workspace_id = report['workspaceId']
        workspace_name =  report['workspace']
        report_id =  report['id']
        report_name =  report['name']
        print(f"Processing report [{report_name}] from workspace [{workspace_name}]")
        # Executing
        create_semantic_model_doc(workspace_name=workspace_name, dataset_name=report_name)

    print("All docx files were generated!")

# Exec
access_token = get_token(tenant_id, client_id, client_secret)
get_tenant_metadata(access_token)
get_info_datasets(cmd, tenant_id, client_id, client_secret, dax_query)
get_dataflows(access_token) 
get_pro_datasets(access_token)
get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop)
create_documentation()
