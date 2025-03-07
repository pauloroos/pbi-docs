# import requests
import os
# import json
import subprocess
import time
import psutil
import pandas as pd
import clr
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from pathlib import Path
import re

pbix_folder = "results/exported_pbix"

# Path DAX Studio CLI
cmd = r"C:\Program Files\DAX Studio\dscmd.exe"

# Path Analysis Services
ssas_dll = r"C:\Program Files\DAX Studio\bin\Microsoft.AnalysisServices.dll"

# Path Power BI Desktop
pbi_desktop = r"C:\Program Files\WindowsApps\Microsoft.MicrosoftPowerBIDesktop_2.140.1351.0_x64__8wekyb3d8bbwe\bin\PBIDesktop.exe"

# DAX Query 
dax_query = """
DEFINE

// Source Tables
VAR __tables = INFO.TABLES()
VAR __columns = INFO.COLUMNS()
VAR __measures = INFO.MEASURES()
VAR __formats = INFO.FORMATSTRINGDEFINITIONS()
VAR __relationships = INFO.RELATIONSHIPS()
VAR __calculationGroups = INFO.CALCULATIONGROUPS()
VAR __calculationItems = INFO.CALCULATIONITEMS()
VAR __dependencies = INFO.CALCDEPENDENCY()
VAR __partitions = INFO.PARTITIONS()


// Measures
VAR __measuresPre = 
	ADDCOLUMNS(
    		__measures,
    		"Table",
		VAR TableID = [TableID] RETURN
	    	MAXX(FILTER(__tables, [ID] = TableID), [Name]),
    		"Format",
		VAR MeasureID = [ID] RETURN
	    	MAXX(FILTER(__formats, [ObjectID] = MeasureID), [Expression])
	)

VAR __measuresResult =
SELECTCOLUMNS(
	__measuresPre,
	"tableName", [Table],
	"name", [Name],
	"expression", [Expression],
	"format", [FormatString],
	"isHidden", [IsHidden],
	"description", [Description],
	"type",
	SWITCH(
		[DataType],
		2, "String",
		6, "Whole Number",
		8, "Double",
		9, "Datetime",
		10, "Currency",
		11, "Boolean"
	),
	"displayFolder", [DisplayFolder]
) 

// Columns
VAR __columnsPre =
ADDCOLUMNS(
	FILTER(
		__columns,
		NOT CONTAINSSTRING([ExplicitName], "RowNumber")
	),
	"TableName",
	VAR TableID = [TableID] RETURN
	MAXX(FILTER(__tables, [ID] = TableID), [Name])
)

VAR __columnsResult =
SELECTCOLUMNS(
	__columnsPre,
    "tableName", [TableName],
    "name", COALESCE([ExplicitName], [InferredName]),
	"column", 
		"'" & [TableName] & "'" & 
		"[" & COALESCE([ExplicitName], [InferredName]) & "]",
    "sortedBy",
		VAR OrderID = [SortByColumnID] RETURN
	    MAXX(
			FILTER(__columnsPre, [ID] = OrderID), 
			"'" & [TableName] & "'" & 
		"[" & COALESCE([ExplicitName], [InferredName]) & "]"),
    "format", COALESCE([FormatString], "String"),
    "displayFolder", [DisplayFolder],
    "isHidden", [IsHidden],
    "expression", [Expression]
)

// Relationships
VAR relationshipsPre =
ADDCOLUMNS(
    __relationships,
    "FromTable",
		VAR FromTable = [FromTableID] RETURN
		MAXX(FILTER(__columnsPre, [TableID] = FromTable), [TableName]),
    "ToTable",
		VAR ToTable = [ToTableID] RETURN
	    MAXX(FILTER(__columnsPre, [TableID] = ToTable), [TableName]),
    "FromColumn",
		VAR FromColumn = [FromColumnID] RETURN
	    MAXX(
			FILTER( __columnsPre, [ID] = FromColumn ),
			COALESCE( [ExplicitName], [InferredName] )),
    "ToColumn",
	VAR ToColumn = [ToColumnID]
	RETURN
	    MAXX(
		FILTER( __columnsPre, [ID] = ToColumn ),
		COALESCE( [ExplicitName], [InferredName] )
	    )
)

VAR __relationshipsResult =
SELECTCOLUMNS(
    relationshipsPre,
    "from", "'" & [FromTable] & "'" & "[" & [FromColumn] & "]",
	"fromCardinality", IF([FromCardinality] = 2, "*", [FromCardinality]),
	"to",	"'" & [ToTable] & "'" & "[" & [ToColumn] & "]",
    "toCardinality", IF( [ToCardinality] = 2, "*", [ToCardinality]),
    "isActive", [IsActive],
    "isBidirectional", IF([CrossFilteringBehavior] = 2, "True", "False"),
    "relationship",
		VAR fromConcat = "'" & [FromTable] & "'" & "[" & [FromColumn] & "]"
		VAR toConcat = "'" & [ToTable] & "'" & "[" & [ToColumn] & "]"
		VAR fromCardinality = IF([FromCardinality] = 2, "*", [FromCardinality])
		VAR toCardinality = IF([ToCardinality] = 2, "*", [ToCardinality])
		VAR arrow = IF([CrossFilteringBehavior] = 2, "<->", "<-")
		RETURN
	    	fromConcat & " " & fromCardinality & " " & arrow & " " & toCardinality & " " & toConcat
)

// Calculation Groups
VAR __calculationItemsPre =
ADDCOLUMNS(
    ADDCOLUMNS(
	__calculationItems,
	"TableID",
	    VAR CalculationGroup_ID = [CalculationGroupID] RETURN
		MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [TableID]),
	"Precedence",
	    VAR CalculationGroup_ID = [CalculationGroupID] RETURN
		MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [Precedence])
    ),
    "CalculationGroup",
		VAR Table_ID = [TableID] RETURN
	    MAXX(FILTER(__tables, [ID] = Table_ID), [Name]),
    "Format",
		VAR Format_ID = [FormatStringDefinitionID] RETURN
	    MAXX(FILTER(__formats, [ID] = Format_ID), [Expression]),
    "CalculationGroupColumn",
		VAR Table_ID = [TableID]
		VAR Coluna_ID = MINX(FILTER(__columns, [TableID] = Table_ID), [ID])
		RETURN
		    MAXX(
				FILTER( __columns, [ID] = Coluna_ID ),
				COALESCE( [ExplicitName], [InferredName] )
		    ),
    "OrdinalColumn",
		VAR Table_ID = [TableID]
		VAR Coluna_ID = MAXX(FILTER(__columns, [TableID] = Table_ID), [ID])
		RETURN
		    MAXX(
				FILTER( __columns, [ID] = Coluna_ID ),
				COALESCE( [ExplicitName], [InferredName] )
		    )
)

VAR __calculationGroupsResult =
SELECTCOLUMNS(
    __calculationItemsPre,
    "calculationGroup", [CalculationGroup],
    "precedence", [Precedence],
    "calculationItem", [Name],
    "expression", [Expression],
    "format", [Format],
    "ordinal", [Ordinal],
    "calculationGroupColumn", [CalculationGroupColumn],
    "ordinalColumn", [OrdinalColumn]
)

VAR __parametersResult =
SELECTCOLUMNS(
	SUMMARIZE(
	    FILTER(
			__dependencies,
			[REFERENCED_OBJECT_TYPE] = "M_EXPRESSION" && 
			CONTAINSSTRING( [REFERENCED_EXPRESSION], "IsParameterQuery" )
		),
		[REFERENCED_OBJECT],
		[REFERENCED_EXPRESSION]
	),
    "name", [REFERENCED_OBJECT],
    "expression", [REFERENCED_EXPRESSION]
)


// Queries M or DAX
VAR __partitionsResult =
SELECTCOLUMNS(
    __partitions,
    "tableName", 
	IF(
        NOT(ISBLANK([Name])) && LEN([Name]) > 37,
        LEFT([Name], LEN([Name]) - 37),
        [Name]
    ),
    "description", [Description],
    "queryDefinition", [QueryDefinition],
    "modifiedTime", [ModifiedTime],
    "refreshedTime", [RefreshedTime],
    "type", SWITCH( [Type], 4, "M", 2, "DAX", 7, "Internal" )
)

// Output
EVALUATE __measuresResult ORDER BY [tableName] ASC, [name] ASC
EVALUATE __columnsResult ORDER BY [tableName] ASC, [name] ASC
EVALUATE __relationshipsResult ORDER BY [relationship] ASC
EVALUATE __calculationGroupsResult ORDER BY [precedence] ASC, [ordinal] ASC
EVALUATE __parametersResult 
EVALUATE __partitionsResult ORDER BY [tableName] ASC 
"""

def get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop):
    
    # Load SSAS assembly
    clr.AddReference(ssas_dll)
    from Microsoft.AnalysisServices import Server 

    query = dax_query

    # Wait for Power BI Open
    def wait_for_powerbi_to_open(timeout=60, check_interval=5):
        """
        Waits for Power BI Desktop to open and start the Analysis Services server.

        timeout: Maximum wait time in seconds.
        check_interval: Interval between checks in seconds.
        """
        elapsed_time = 0
        while elapsed_time < timeout:
            server_name = get_powerbi_port()
            if server_name != "No instance of Power BI Desktop found.":
                print(f"Power BI Desktop started! Server found at: {server_name}")
                return server_name
            print(f"Waiting for Power BI Desktop to open... ({elapsed_time}s)")
            time.sleep(check_interval)
            elapsed_time += check_interval

        print("Timeout reached. Power BI Desktop did not start correctly.")
        return None

    def get_powerbi_port():
        # Search the process msmdsrv.exe (Analysis Services Power BI)
        for proc in psutil.process_iter(attrs=['pid', 'name']):
            if proc.info['name'] == "msmdsrv.exe":
                pid = proc.info['pid']
                
                # Run netstat to capture active conections
                result = subprocess.run(["netstat", "-ano"], capture_output=True, text=True)
                connections = result.stdout.splitlines()
                
                # Filter corresponding line to PID of msmdsrv.exe
                for line in connections:
                    if str(pid) in line and "LISTENING" in line:
                        match = re.search(r':(\d+)', line)
                        if match:
                            return f"localhost:{match.group(1)}"
        
        return "No instance of Power BI Desktop found."

    # Function to get the database ID
    def get_ssas_database_id(server_name):
        
        server = Server()
        try:
            server.Connect(f"Data Source={server_name}")
            # Assuming there is only one active database
            if server.Databases.Count > 0:
                db_id = server.Databases[0].ID
                return db_id
            else:
                return None
        finally:
            server.Disconnect()

    def wait_for_powerbi_to_close(timeout=30, check_interval=2):
        """
        Waits for Power BI Desktop to close completely.
        """
        elapsed_time = 0
        while elapsed_time < timeout:
            if "PBIDesktop.exe" not in (p.name() for p in psutil.process_iter()):
                print("Power BI Desktop closed successfully.")
                return True
            print(f"Waiting for Power BI Desktop to close... ({elapsed_time}s)")
            time.sleep(check_interval)
            elapsed_time += check_interval
        print("Timeout waiting for Power BI Desktop to close.")
        return False

    # Iterate through each workspace
    for pbix_file_name in pbix_files_names:
        print(f"Processing: {pbix_file_name}.pbix")
        pbix_path = os.path.abspath(f"results/exported_pbix/{pbix_file_name}.pbix")

        # Execute
        if not os.path.exists(pbix_path):
            print(f"Error: The file {pbix_file_name}.pbix was not found.")
        else:
            try:
                subprocess.Popen([pbi_desktop, pbix_path])
                print(f"Power BI Desktop opened the file {pbix_file_name}.pbix.")
            except Exception as e:
                print(f"Error opening Power BI Desktop: {e}")

        # Get Power BI Desktop Server
        server_name = wait_for_powerbi_to_open()

        if not server_name:
            print("Failed to start Power BI. Retrying...")
            time.sleep(5)
            server_name = wait_for_powerbi_to_open()

        if server_name:
            print(f"Server: {server_name}")
            database_id = get_ssas_database_id(server_name=server_name)
            if database_id:
                print(f"Database: {database_id}")
            else:
                print("No database found. Retrying after waiting...")
                time.sleep(5)
                database_id = get_ssas_database_id(server_name=server_name)
                if not database_id:
                    print("Error: Unable to retrieve database ID.")
                    continue  # Skip this, try next
        else:
            print("No server found. Skipping this report.")
            continue

        # Dax Studio Variables
        server = server_name
        database = database_id

        # Output path
        output_dir = f"results/datasets_info/{pbix_file_name}"
        os.makedirs(output_dir, exist_ok=True)
        output_path = f"{output_dir}/extract.csv"

        # Clean up any existing extract files before execution
        for old_name in ["extract.csv", "extract_2.csv", "extract_3.csv", 
                        "extract_4.csv", "extract_5.csv", "extract_6.csv"]:
            old_path = os.path.join(output_dir, old_name)
            if os.path.exists(old_path):
                os.remove(old_path)
                print(f"    Removed existing source file: {old_name}")

        max_retries = 3
        retry_delay = 5

        for attempt in range(max_retries):
            try:
                # Executing
                subprocess.run([
                    cmd, "csv", output_path,
                    "-s", server,
                    "-d", database,
                    "-q", query
                ], check=True)
                break
            except subprocess.CalledProcessError:
                print(f"Error executing dscmd. Attempt {attempt + 1} of {max_retries}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    print("Persistent failure executing dscmd. Skipping this report.")


        # Close Power BI Desktop
        subprocess.run(["taskkill", "/F", "/IM", "PBIDesktop.exe"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print("Power BI Desktop was closed.")
        
        
        # Remaping the extracted files
        rename_map = {
            "extract.csv": "measures.csv",
            "extract_2.csv": "columns.csv",
            "extract_3.csv": "relationships.csv",
            "extract_4.csv": "calculation_groups.csv",
            "extract_5.csv": "parameters.csv",
            "extract_6.csv": "partitions.csv"
        }

        # Rename files
        for old_name, new_name in rename_map.items():
            old_path = os.path.join(output_dir, old_name)
            new_path = os.path.join(output_dir, new_name)

            # Check if old file exists
            if os.path.exists(old_path):
                # If new file already exists, remove it first
                if os.path.exists(new_path):
                    os.remove(new_path)
                    print(f"    Removed existing file: {new_name}")
            
                os.rename(old_path, new_path)
                print(f"    Renamed: {old_name} -> {new_name}")
            else:
                print(f"    File not found: {old_name}")

        
        # Wait for the process to actually terminate before opening the next file
        time.sleep(5)
        wait_for_powerbi_to_close()


    print("Processing completed all pbix files in local folder.")

def create_documentation():
    def create_semantic_model_doc(dataset_name):
        
        directory = f"results/datasets_info/{dataset_name}"
        doc = Document()
        
        # Ajustar as margens do documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # Custom Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(3)

        # Main Title
        doc.add_heading("Semantic Model Technical Documentation", level=1)
        
        # Overview
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(f"Semantic Model: {dataset_name}")
        doc.add_paragraph("This document describes the extracted semantic model structure, including tables, columns, relationships, partitions, measures, calculation groups, and parameters.")
        
        # Load csv files
        file_paths = {
            "Calculation Groups": os.path.join(directory, "calculation_groups.csv"),
            "Columns": os.path.join(directory, "columns.csv"),
            "Measures": os.path.join(directory, "measures.csv"),
            "Parameters": os.path.join(directory, "parameters.csv"),
            "Partitions": os.path.join(directory, "partitions.csv"),
            "Relationships": os.path.join(directory, "relationships.csv"),
        }
        
        dataframes = {name: pd.read_csv(path, encoding="utf-8", delimiter=";") for name, path in file_paths.items() if os.path.exists(path)}

        # Nan to ""
        for key in dataframes:
            dataframes[key] = dataframes[key].fillna("")

        # Apply style tables
        def apply_table_style(table):
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.paragraphs:
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i % 2 == 0:
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F3F3F3"/>'.format(nsdecls('w'))))

        # Add sections
        sections = {
            "Columns": ["Columns", ["Table", "Column", "Sorted By", "Format", "Display Folder", "Hidden", "Expression"], ["tableName","name","sortedBy","format","displayFolder","isHidden","expression"]],
            "Partitions": ["Partitions", ["Table", "Type", "Query"], ["tableName", "type", "queryDefinition"]],
            "Relationships": ["Relationships", ["Relationship", "isActive", "isBidirectional"], ["relationship", "isActive", "isBidirectional"]],
            "Measures": ["Measures", ["Name", "Expression", "Description", "Format"], ["name", "expression", "description", "format"]],
            "Calculation Groups": ["Calculation Groups", ["Group", "Item", "Expression", "Ordinal", "Format"], ["calculationGroup", "calculationItem", "expression", "ordinal", "format"]],
            "Parameters": ["Parameters", ["Name", "Expression"], ["name", "expression"]],
        }
        
        for key, (title, headers, columns) in sections.items():
            if key in dataframes:
                doc.add_heading(title, level=2)
                table = doc.add_table(rows=1, cols=len(headers))

                # Set column widths

                if key == "Partitions":
                    table.columns[0].width = Cm(1)  # Table
                    table.columns[1].width = Cm(1)  # Type
                    table.columns[2].width = Cm(48)  # Query

                elif key == "Relationships":
                    table.columns[0].width = Cm(24)  # Relationship
                    table.columns[1].width = Cm(1)  # isActive
                    table.columns[2].width = Cm(1)  # isBidirectional

                elif key == "Measures":
                    table.columns[0].width = Cm(2)  # Name
                    table.columns[1].width = Cm(36)  # Expression
                    table.columns[2].width = Cm(36)  # Description
                    table.columns[3].width = Cm(2)  # Format

                elif key == "Calculation Groups":
                    table.columns[0].width = Cm(2)  # Group
                    table.columns[1].width = Cm(2)  # Item
                    table.columns[2].width = Cm(36)  # Expression
                    table.columns[3].width = Cm(2)  # Ordinal
                    table.columns[4].width = Cm(8)  # Format

                elif key == "Parameters":
                    table.columns[0].width = Cm(2)  # Name
                    table.columns[1].width = Cm(36)  # Expression

                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for _, row in dataframes[key].iterrows():
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        if i < len(row_cells):
                            row_cells[i].text = str(row[col]) if col in row else ""
                
                apply_table_style(table)
        
        doc.add_paragraph("---")
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Document generated on: {now}")
        
        # Ensure the directory exists
        dir_path = "results/documentation"
        os.makedirs(dir_path, exist_ok=True)
        
        file_path = f"{dir_path}/{dataset_name}.docx"
        doc.save(file_path)
        print(f"File {file_path} saved.")

    for pbix_file_name in pbix_files_names:
        print(f"Processing report {pbix_file_name}")
        create_semantic_model_doc(dataset_name=pbix_file_name)

    print("All docx files were generated!")

# Exec
pbix_files_names = [os.path.splitext(f)[0] for f in os.listdir(pbix_folder) if f.endswith('.pbix')]  
get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop)
create_documentation()
